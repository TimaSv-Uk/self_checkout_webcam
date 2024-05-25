from kivymd.app import MDApp
from kivymd.uix.widget import MDWidget
from kivymd.uix.datatables import MDDataTable
from kivy.metrics import dp
from kivymd.uix.snackbar import MDSnackbar
from kivymd.uix.label import MDLabel
from kivymd.uix.boxlayout import MDBoxLayout
from kivymd.uix.button import MDFillRoundFlatButton

import pyodbc
import os
from docx import Document
from pandas import DataFrame
from docx2pdf import convert
from plyer import filechooser


class Analytics(MDWidget):
    data_tables = {}
    data_table_save = {}

    def open_table(self, table_name: str, description: str):
        if table_name not in self.data_tables.keys():
            self.add_data(table_name, description)
        else:
            self.ids.analytics_layout.remove_widget(self.data_table_save[table_name])
            self.ids.analytics_layout.remove_widget(self.data_tables[table_name])
            self.add_data(table_name, description)

    def add_data(self, table_name: str, description: str):
        products = MDApp.get_running_app().select_as_dict(
            f"select * from {table_name};"
        )

        self.data_tables[table_name] = MDDataTable(
            use_pagination=True,
            size_hint=(1, 0.7),
            column_data=[(col_name, dp(70)) for col_name in products[0].keys()],
            row_data=[(list(pr.values())) for pr in products],
        )
        self.data_table_save[table_name] = MDBoxLayout(
            orientation="horizontal",
            size_hint_y=None,
            size=(200, 50),
        )
        self.data_table_save[table_name].add_widget(
            MDLabel(text=description, font_style="H6")
        )

        self.data_table_save[table_name].add_widget(
            MDFillRoundFlatButton(
                text="Зберегти",
                on_press=lambda instance: self.save_data_to_file(
                    table_name, description
                ),
            )
        )
        self.ids.analytics_layout.add_widget(self.data_table_save[table_name])
        self.ids.analytics_layout.add_widget(self.data_tables[table_name])

    def save_data_to_file(self, table_name, title_in_file):
        path = filechooser.save_file(
            title="Save File As",
            filters=[
                ("Word Files", "*.docx"),
                ("Excel Files", "*.xlsx"),
                ("PDF Files", "*.pdf"),
            ],
        )
        if not path:
            print("User cancelled file selection.")
            return
        path = path[0]
        file_type = (
            "excel"
            if path.endswith(".xlsx")
            else "word"
            if path.endswith(".docx")
            else "pdf"
            if path.endswith(".pdf")
            else None
        )
        if file_type is None:
            MDSnackbar(
                MDLabel(
                    text="Недійсний тип файлу. Виберіть «excel», «word» або «pdf».",
                ),
            ).open()
            return
        df = DataFrame(
            self.data_tables[table_name].row_data,
            columns=[col[0] for col in self.data_tables[table_name].column_data],
        )
        try:
            if file_type == "excel":
                df.to_excel(path, index=False)
                os.startfile(path)
            elif file_type in ["word", "pdf"]:
                doc = Document()
                doc.add_heading(table_name, 1)
                doc.add_heading(title_in_file, 0)
                table = doc.add_table(df.shape[0] + 1, df.shape[1])

                for j in range(df.shape[-1]):
                    table.cell(0, j).text = df.columns[j]

                for i in range(df.shape[0]):
                    for j in range(df.shape[-1]):
                        table.cell(i + 1, j).text = str(df.values[i, j])

                word_path = (
                    path if file_type == "word" else path.replace(".pdf", ".docx")
                )
                doc.save(word_path)

                if file_type == "pdf":
                    convert(word_path, path)
                    os.remove(word_path)  # remove the Word file after conversion

                os.startfile(path)
        except Exception:
            MDSnackbar(
                MDLabel(
                    text="Не може зберегти файл",
                ),
            ).open()
            return
